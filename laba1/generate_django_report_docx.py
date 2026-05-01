from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_default_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)


def add_paragraph(
    document: Document,
    text: str,
    *,
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent_cm: float | None = 1.25,
    spacing_after_pt: float = 6,
    spacing_before_pt: float = 0,
    line_spacing: float = 1.5,
) -> None:
    p = document.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(spacing_after_pt)
    p.paragraph_format.space_before = Pt(spacing_before_pt)
    p.paragraph_format.line_spacing = line_spacing
    if first_line_indent_cm is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent_cm)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)


def add_center_line(document: Document, text: str, *, bold: bool = False, size: int = 14) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)


def remove_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "nil")
        tbl_borders.append(elem)
    tbl_pr.append(tbl_borders)


def add_figure(document: Document, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        add_paragraph(document, f"[Пропущено зображення: {image_path.name}]", align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent_cm=None)
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(16))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.space_after = Pt(8)
    cap.paragraph_format.line_spacing = 1.2
    cap_run = cap.add_run(caption)
    cap_run.italic = True
    cap_run.font.name = "Times New Roman"
    cap_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    cap_run.font.size = Pt(13)


def main() -> None:
    repo_root = Path(__file__).resolve().parent.parent
    screenshots_dir = repo_root / "tmp" / "docs" / "screenshots"
    output_dir = repo_root / "output" / "doc"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_file = output_dir / "Django_Lab1_Report.docx"

    doc = Document()
    set_default_style(doc)

    # Title page
    add_center_line(doc, "МІНІСТЕРСТВО ОСВІТИ І НАУКИ УКРАЇНИ", bold=True)
    add_center_line(doc, "ХАРКІВСЬКИЙ НАЦІОНАЛЬНИЙ УНІВЕРСИТЕТ РАДІОЕЛЕКТРОНІКИ", bold=True)
    add_center_line(doc, "Кафедра «Програмна інженерія»")

    doc.add_paragraph("\n")
    doc.add_paragraph("\n")
    add_center_line(doc, "ЗВІТ", bold=True, size=16)
    add_center_line(doc, "з лабораторної роботи", size=14)
    add_center_line(doc, "з дисципліни «Високорівневі мови програмування та фреймворки»", size=14)
    add_center_line(doc, "З лабораторної роботи №1", size=14)
    add_center_line(doc, "Тема: Розробка веб-додатка на Django", size=14)

    doc.add_paragraph("\n")
    doc.add_paragraph("\n")

    table = doc.add_table(rows=3, cols=2)
    table.columns[0].width = Cm(8)
    table.columns[1].width = Cm(8)
    table.cell(0, 0).text = "Виконали:\nст. гр. __________\n__________________"
    table.cell(0, 1).text = "Прийняв:\nасистент кафедри ПІ\n__________________"
    table.cell(1, 0).text = ""
    table.cell(1, 1).text = ""
    table.cell(2, 0).text = ""
    table.cell(2, 1).text = ""

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.line_spacing = 1.2
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(14)

    remove_table_borders(table)

    doc.add_paragraph("\n")
    doc.add_paragraph("\n")
    add_center_line(doc, f"Харків {datetime.now().year}", size=14)

    doc.add_page_break()

    # Main content
    add_paragraph(doc, "Мета роботи", bold=True, first_line_indent_cm=None, align=WD_ALIGN_PARAGRAPH.LEFT, spacing_after_pt=4)
    add_paragraph(
        doc,
        "Ознайомитися з основами розробки веб-додатків на мові Python з використанням фреймворку Django. "
        "Вивчити архітектурний патерн MVC та особливості його реалізації в Django для розділення логіки даних, "
        "інтерфейсу та управління. Навчитися створювати проєкти та додатки, проєктувати моделі даних за допомогою "
        "Django ORM, розробляти контролери (views) для обробки HTTP-запитів та створювати динамічні сторінки "
        "за допомогою системи шаблонів.",
    )

    add_paragraph(doc, "Хід роботи", bold=True, first_line_indent_cm=None, align=WD_ALIGN_PARAGRAPH.LEFT, spacing_after_pt=4)

    add_paragraph(doc, "1. Налаштування проєкту та створення структури", bold=True)
    add_paragraph(
        doc,
        "На першому етапі було підготовлено робоче середовище та структуру Django-проєкту PythonApplication1. "
        "У корені застосунку створено конфігураційний модуль, керуючий файл manage.py, окремий додаток articles, "
        "а також базові файли моделей, представлень, маршрутів і шаблонів. Структуру проєкту приведено до формату, "
        "придатного для подальшого розширення лабораторної роботи.",
    )
    add_figure(doc, screenshots_dir / "fig01_home.png", "Рисунок 1 - Головна сторінка веб-додатка")

    add_paragraph(doc, "2. Створення суперкористувача (Admin User)", bold=True)
    add_paragraph(
        doc,
        "Для керування контентом через вбудований інтерфейс адміністратора було створено обліковий запис "
        "суперкористувача. Авторизація в адмін-панелі забезпечила повний доступ до об'єктів моделі Article "
        "та можливість виконувати CRUD-операції без додаткового коду інтерфейсу на ранньому етапі розробки.",
    )
    add_figure(doc, screenshots_dir / "fig06_admin_login.png", "Рисунок 2 - Авторизація в адмін-панелі Django")

    add_paragraph(doc, "3. Робота з моделями даних (Models)", bold=True)
    add_paragraph(
        doc,
        "У файлі models.py було реалізовано модель Article з полями title, text і published_at. "
        "Для коректного виводу в адмінці визначено метод __str__, а також метадані сортування. "
        "Підхід на основі Django ORM дозволив керувати структурою таблиць і запитами до бази даних "
        "через Python-об'єкти без ручного SQL.",
    )

    add_paragraph(doc, "4. Налаштування та активація бази даних", bold=True)
    add_paragraph(
        doc,
        "В якості СУБД використано SQLite3. Структуру бази активовано через міграції "
        "python manage.py makemigrations та python manage.py migrate. Після цього система "
        "зберігання даних стала готовою до наповнення реальними записами статей.",
    )
    add_figure(doc, screenshots_dir / "fig07_admin_dashboard.png", "Рисунок 3 - Головна сторінка адмін-панелі після входу")

    add_paragraph(doc, "5. Перевірка панелі адміністратора та наповнення бази даних", bold=True)
    add_paragraph(
        doc,
        "Після активації моделей перевірено працездатність адмін-інтерфейсу. "
        "Реалізовано перегляд списку статей, форму створення нової статті та форму редагування. "
        "Базу даних заповнено тестовими записами для подальшого тестування пошуку і відображення.",
    )
    add_figure(doc, screenshots_dir / "fig08_admin_article_list.png", "Рисунок 4 - Перелік статей в адмін-панелі")
    add_figure(doc, screenshots_dir / "fig09_admin_add_form.png", "Рисунок 5 - Форма додавання нової статті")
    add_figure(doc, screenshots_dir / "fig10_admin_edit_form.png", "Рисунок 6 - Форма редагування статті")

    add_paragraph(doc, "6. Розробка представлення (views.py)", bold=True)
    add_paragraph(
        doc,
        "Для передачі даних із бази на веб-сторінку реалізовано функцію article_list. "
        "Вона отримує колекцію об'єктів Article, обробляє параметр пошуку q, формує контекст "
        "і передає результати в шаблон article_list.html. Окремо реалізовано функцію article_detail "
        "для відкриття повного тексту обраної статті.",
    )

    add_paragraph(doc, "7. Налаштування маршрутизації (urls.py)", bold=True)
    add_paragraph(
        doc,
        "У модулі маршрутизації додано URL для головної сторінки списку статей, "
        "детальної сторінки article/<id>/ та стандартної адмін-панелі admin/. "
        "Це забезпечило зв'язок між HTTP-запитами користувача та відповідними обробниками.",
    )

    add_paragraph(doc, "8. Створення HTML-шаблону відображення", bold=True)
    add_paragraph(
        doc,
        "Шаблон article_list.html реалізовано на Django Template Language. "
        "Інтерфейс містить форму пошуку, блок статистики результатів, картки статей, "
        "переходи на детальний перегляд та посилання на адмін-панель.",
    )
    add_figure(doc, screenshots_dir / "fig02_search_input.png", "Рисунок 7 - Форма введення ключових слів на головній сторінці")

    add_paragraph(doc, "9. Реалізація пошуку статей за ключовими словами", bold=True)
    add_paragraph(
        doc,
        "Пошук реалізовано через ORM-фільтрацію з використанням Q-виразів по полях title і text. "
        "Запит користувача розбивається на ключові слова, після чого формується об'єднаний фільтр "
        "за принципом логічного OR. Додано обробку випадку, коли результати відсутні.",
    )
    add_figure(doc, screenshots_dir / "fig03_search_results.png", "Рисунок 8 - Результати пошуку за ключовими словами")
    add_figure(doc, screenshots_dir / "fig04_no_results.png", "Рисунок 9 - Відображення стану «збігів не знайдено»")

    add_paragraph(doc, "10. Створення методу детального перегляду статті (views.py)", bold=True)
    add_paragraph(
        doc,
        "Для перегляду повного контенту реалізовано окремий маршрут та шаблон article_detail.html. "
        "Сторінка відображає дату, заголовок, повний текст, кнопку повернення до списку та перехід до редагування "
        "в адмін-панелі. Отримання конкретної статті виконується через get_object_or_404.",
    )
    add_figure(doc, screenshots_dir / "fig05_article_detail.png", "Рисунок 10 - Сторінка детального перегляду статті")

    add_paragraph(doc, "11. Створення візуальної розмітки та адаптивності інтерфейсу", bold=True)
    add_paragraph(
        doc,
        "Для поліпшення зовнішнього вигляду застосовано статичні CSS-файли, градієнтні фони, "
        "карткову структуру контенту, а також SVG-анімації в hero-блоці. "
        "Додатково реалізовано адаптацію інтерфейсу під мобільні пристрої через media queries.",
    )
    add_figure(doc, screenshots_dir / "fig11_admin_search.png", "Рисунок 11 - Пошук статей в адмін-панелі")
    add_figure(doc, screenshots_dir / "fig12_mobile_view.png", "Рисунок 12 - Адаптивне відображення сайту на мобільному екрані")

    add_paragraph(doc, "Висновки", bold=True, first_line_indent_cm=None, align=WD_ALIGN_PARAGRAPH.LEFT, spacing_after_pt=4)
    add_paragraph(
        doc,
        "Під час виконання лабораторної роботи було повністю реалізовано веб-додаток на Django "
        "для керування статтями. У процесі роботи відпрацьовано створення моделей та міграцій, "
        "адміністративне керування контентом, реалізацію публічного інтерфейсу, пошук за ключовими словами "
        "та сторінку детального перегляду. Додатково виконано стилізацію інтерфейсу через статичні файли "
        "і впроваджено анімовані SVG-елементи, що підвищило якість користувацького досвіду.",
    )

    doc.add_page_break()
    add_paragraph(doc, "Контрольні запитання", bold=True, first_line_indent_cm=None, align=WD_ALIGN_PARAGRAPH.LEFT, spacing_after_pt=6)

    qa = [
        "1. Переваги Django: висока швидкість розробки, принцип DRY, вбудована адмін-панель, високий рівень безпеки та велика екосистема бібліотек.",
        "2. Структура проєкту: конфігураційний модуль (settings, urls, wsgi/asgi) та один або кілька додатків із власними моделями, views, templates і static.",
        "3. Основні компоненти Django: моделі, URL-маршрути, представлення, шаблони та middleware.",
        "4. Налаштування і запуск: встановити Python і Django, створити проєкт, виконати міграції та запустити сервер командою python manage.py runserver.",
        "5. Розгортання в продакшен: DEBUG=False, ALLOWED_HOSTS, окремий веб-сервер, WSGI/ASGI-процес та надійна база даних.",
        "6. Створення моделі даних: у файлі models.py визначається клас, успадкований від models.Model, із потрібними полями.",
        "7. Можливості ORM: створення, читання, оновлення та видалення даних через Python-інтерфейс, без прямого SQL.",
        "8. Використання ORM: через менеджер Model.objects (all, filter, create, get, update, delete).",
        "9. URL для різних запитів: задаються у urls.py, а метод запиту (GET/POST/PUT/DELETE) обробляється у view.",
        "10. Динамічні сторінки: створюються через DTL із тегами {% %} та змінними {{ }}.",
        "11. Клас Form: спрощує створення HTML-форм, валідацію введення і перетворення даних.",
        "12. Сторонні пакети: встановлюються через pip і підключаються в налаштуваннях проєкту.",
        "13. Middleware: прошарок між запитом і відповіддю для обробки сесій, безпеки, кешування тощо.",
        "14. Інструменти тестування: TestCase, test client, ізольована тестова база та запуск через manage.py test.",
        "15. Автентифікація та авторизація: вбудований модуль django.contrib.auth з користувачами, групами і правами доступу.",
        "16. Кешування: підтримується на рівні сторінок, шаблонних фрагментів і запитів через Redis/Memcached/DB.",
        "17. Система адміністрування: моделі реєструються у admin.py, після чого доступні в /admin/.",
        "18. Створення API: найчастіше використовують Django REST Framework із серіалізаторами та viewsets.",
    ]
    for item in qa:
        add_paragraph(doc, item)

    doc.save(output_file)
    print(f"Generated: {output_file}")


if __name__ == "__main__":
    main()
